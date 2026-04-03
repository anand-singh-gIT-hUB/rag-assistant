"""
app/ingestion/loaders/docx_loader.py
"""
from pathlib import Path

from docx import Document

from app.ingestion.loaders.base import BaseLoader, PageContent
from app.core.exceptions import DocumentProcessingError


class DocxLoader(BaseLoader):
    supported_extensions = {".docx"}

    def load(self, path: Path) -> list[PageContent]:
        try:
            doc = Document(str(path))
            pages: list[PageContent] = []
            current_heading: str | None = None
            buffer: list[str] = []

            for para in doc.paragraphs:
                style = para.style.name.lower()
                text = para.text.strip()
                if not text:
                    continue
                if style.startswith("heading"):
                    if buffer:
                        pages.append(
                            PageContent(
                                text="\n".join(buffer),
                                section_title=current_heading,
                            )
                        )
                        buffer = []
                    current_heading = text
                else:
                    buffer.append(text)

            if buffer:
                pages.append(
                    PageContent(text="\n".join(buffer), section_title=current_heading)
                )

            return pages or [PageContent(text="\n".join(p.text for p in doc.paragraphs if p.text.strip()))]
        except Exception as e:
            raise DocumentProcessingError(
                f"Failed to parse DOCX: {path.name}", detail=str(e)
            ) from e
